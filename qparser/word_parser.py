import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.text.paragraph import Paragraph

from .base_parser import BaseParser


class WordParser(BaseParser):

    def __init__(self, add_alphabet: bool = False, *args, **kwargs):
        """
        :param add_alphabet: add alphabet enumeration to answers if true (default=False)
        """
        self.add_alphabet = add_alphabet
        super().__init__(*args, **kwargs)

    def _parse_question(self, raw_question: list[Paragraph], chapter: str) \
            -> tuple[str, list[str], list[str], str]:
        """ Parse question from list of it's strings

        :param raw_question: list of strings of question and answers in docx.Paragraph type
        :param chapter: name of question chapter
        :return: tuple(question, answers, correct_answer, chapter)
        """

        question = self._refactor_question(raw_question[0].text)
        answers, correct_answer = self._refactor_answers(raw_question[1:])

        return question, answers, correct_answer, chapter

    @classmethod
    def _refactor_question(cls, question: str) -> str:
        """Removes start numeration, any symbols like + and so on.
        Also changes first symbol to uppercase

        :param question: str with raw question
        :return: refactored question
        """

        result = re.sub(r'^[+]?\d+\s?[.]', '', question.strip()).strip()  # Removes start numeration and so on
        return result[0].upper() + result[1:]  # Changes first symbol to uppercase

    def _refactor_answers(self, raw_answers: list[Paragraph]) -> tuple[list[str], list[str]]:
        """Removes numeration of answers and last symbol (like ';', '.' and etc).
        Also finds correct answers

        :param raw_answers: list of raw answers in docx format
        :return: tuple(answers, correct) -> where answers - list of answers
        and correct - list of correct answers
        """

        result = []
        correct = []
        alphabet = self._alphabet_enum(lang_start_symbol='а')
        for number, answer in enumerate(raw_answers):
            cur_string = re.sub(r'^\w[).]\s?[.]?', '', answer.text.strip()).strip()  # Remove numeration of answers
            cur_string = re.sub(r'[.;+]$', '', cur_string.strip()).strip()  # Remove end symbols (like '.', ';', '+')
            if self.add_alphabet:
                result.append(alphabet[number] + ') ' + cur_string)
            else:
                result.append(cur_string)

            for symbol in answer.runs:  # find correct
                if symbol.bold:
                    if self.add_alphabet:
                        correct.append(alphabet[number])
                    else:
                        correct.append(str(number))
                    break
        return result, correct

    @classmethod
    def _alphabet_enum(cls, lang_start_symbol: str = 'а') -> str:
        """Creates a str of alphabet, like: 'abcde..ABC..Z'.

        :param lang_start_symbol: start alphabet symbol of resulting lang
        :return: alphabet string
        """
        start_ord = ord(lang_start_symbol)
        alphabet = ''.join([chr(i) for i in range(start_ord, start_ord + 32)])
        return alphabet

    @classmethod
    def _document_to_raw_questions(cls, document: Document) -> tuple[list[list[Paragraph]], str]:
        """Parses document to list of raw_question, which in contains: question and answers,
        and chapter name

        :param document: opened document in docx.Document format
        :return: tuple(raw_questions, chapter)
        """

        chapter = document.paragraphs[0].text.strip()
        paragraphs = document.paragraphs[2:]  # skip chapter name

        raw_questions = []
        raw_question = []
        for paragraph in paragraphs:
            if paragraph.text.strip():  # if not empty string
                raw_question.append(paragraph)
            else:
                raw_questions.append(raw_question[:])
                raw_question.clear()
        return raw_questions, chapter

    def parse(self, path: str) -> pd.DataFrame:
        """Parses word document of questionnaires to pandas.DataFrame format

        :param path: path to word document
        :return: DataFrame with columns=['Question', 'Answers', 'Correct', 'Chapter']
        """
        document = Document(path)

        raw_questions, chapter = self._document_to_raw_questions(document)
        questions = [self._parse_question(raw, chapter) for raw in raw_questions]

        return pd.DataFrame(questions, columns=['Question', 'Answers', 'Correct', 'Chapter'])
